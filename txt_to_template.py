from docx import Document
import re

def generate_output():
    def extract_style(run):
        style = {
            'bold': run.bold,
            'italic': run.italic,
            'font_name':run.font.name,
            'color':run.font.color.rgb
        }
        return style

    def fill_template(template_path, mcq_file, owq_file):
        # Load the template document
        doc = Document(template_path)
        style = doc.styles['Normal']
        style.font.name = 'Tahoma'

        # Read multiple choice questions from file
        with open(mcq_file, 'r') as mcq_file:
            mcq_data = mcq_file.read()
        
        n = len(re.findall(r'Q(?:1[0-9]|20|[1-9])\.', mcq_data))
        
        # Fill in placeholders for multiple choice questions
        for i in range(1, n+1):
            # Parse question, answer options, and correct answer from each entry
            question_start_index = mcq_data.find(f"Q{i}.")+4
            question_end_index = mcq_data.find("\n", question_start_index)
            question = mcq_data[question_start_index:question_end_index].strip()
        
            answer_option_a_start_index = question_end_index + 4
            answer_option_a_end_index = mcq_data.find("\n", answer_option_a_start_index)
            answer_option_a_text = mcq_data[answer_option_a_start_index:answer_option_a_end_index].strip()

            answer_option_b_start_index = answer_option_a_end_index + 4
            answer_option_b_end_index = mcq_data.find("\n", answer_option_b_start_index)
            answer_option_b_text = mcq_data[answer_option_b_start_index:answer_option_b_end_index].strip()
            
            answer_option_c_start_index = answer_option_b_end_index + 4
            answer_option_c_end_index = mcq_data.find("\n", answer_option_c_start_index)
            answer_option_c_text = mcq_data[answer_option_c_start_index:answer_option_c_end_index].strip()

            answer_option_d_start_index = answer_option_c_end_index + 4
            answer_option_d_end_index = mcq_data.find("\n", answer_option_d_start_index)
            answer_option_d_text = mcq_data[answer_option_d_start_index:answer_option_d_end_index].strip()

            correct_answer_start_index = mcq_data.find("Answer:", answer_option_d_end_index) + len("Answer:")
            correct_answer_end_index = mcq_data.find("\n", correct_answer_start_index)
            correct_answer = mcq_data[correct_answer_start_index:correct_answer_end_index].strip()

            placeholder_question = f'{{{{MCQ{i}}}}}'
            placeholder_answer_a = f'{{{{MCQ{i}_a}}}}'
            placeholder_answer_b = f'{{{{MCQ{i}_b}}}}'
            placeholder_answer_c = f'{{{{MCQ{i}_c}}}}'
            placeholder_answer_d = f'{{{{MCQ{i}_d}}}}'
            placeholder_correct = f'{{{{MCQ{i}_correct}}}}'

            for p in doc.paragraphs:
                if placeholder_question in p.text:                
                    style = extract_style(p.runs[0])
                    p.text = ''
                    run = p.add_run(f'Q{i}. {question}')
                    run.bold = style['bold']
                    run.italic = style['italic']
                    run.font.name = style['font_name']
                    run.font.color.rgb = style['color']

                if placeholder_answer_a in p.text:
                    style = extract_style(p.runs[0])
                    p.text = ''
                    run = p.add_run(f'a) {answer_option_a_text}')
                    run.bold = style['bold']
                    run.italic = style['italic']
                    run.font.name = style['font_name']
                    run.font.color.rgb = style['color']
                
                if placeholder_answer_b in p.text:
                    style = extract_style(p.runs[0])
                    p.text = ''
                    run = p.add_run(f'b) {answer_option_b_text}')
                    run.bold = style['bold']
                    run.italic = style['italic']
                    run.font.name = style['font_name']
                    run.font.color.rgb = style['color']

                if placeholder_answer_c in p.text:
                    style = extract_style(p.runs[0])
                    p.text = ''
                    run = p.add_run(f'c) {answer_option_c_text}')
                    run.bold = style['bold']
                    run.italic = style['italic']
                    run.font.name = style['font_name']
                    run.font.color.rgb = style['color']

                if placeholder_answer_d in p.text:
                    style = extract_style(p.runs[0])
                    p.text = ''
                    run = p.add_run(f'd) {answer_option_d_text}')
                    run.bold = style['bold']
                    run.italic = style['italic']
                    run.font.name = style['font_name']
                    run.font.color.rgb = style['color']

                if placeholder_correct in p.text:
                    style = extract_style(p.runs[0])
                    p.text = ''
                    run = p.add_run(f'Answer: {correct_answer}')
                    run.bold = style['bold']
                    run.italic = style['italic']
                    run.font.name = style['font_name']
                    run.font.color.rgb = style['color']
                
        print("MCQ part completed...")

        # Read multiple choice questions from file
        with open(owq_file, 'r') as owq_file:
            owq_data = owq_file.read()
        
        m = len(re.findall(r'Q(?:1[0-9]|20|[1-9])\.', owq_data))

        # Fill in placeholders for open-written questions
        for i in range(1, m+1):
            # Parse question and correct answer from each entry
            question_start_index = owq_data.find(f"Q{i}.")+4
            question_end_index = owq_data.find("\n", question_start_index)
            question = owq_data[question_start_index:question_end_index].strip()
        
            correct_answer_start_index = question_end_index+1
            correct_answer_end_index = owq_data.find("\n", correct_answer_start_index)
            correct_answer = owq_data[correct_answer_start_index:correct_answer_end_index].strip()

            placeholder_question = f'{{{{OWQ{i}}}}}'
            placeholder_correct = f'{{{{OWQ{i}_correct}}}}'

            for p in doc.paragraphs:
                if placeholder_question in p.text:                
                    style = extract_style(p.runs[0])
                    p.text = ''
                    run = p.add_run(f'Q{i}. {question} (5 marks)')
                    run.bold = style['bold']
                    run.italic = style['italic']
                    run.font.name = style['font_name']
                    run.font.color.rgb = style['color']
                    
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if placeholder_correct in paragraph.text:
                                    style = extract_style(paragraph.runs[0])
                                    paragraph.text = ''
                                    run = paragraph.add_run(f'Answer: {correct_answer}')
                                    run.bold = style['bold']
                                    run.italic = style['italic']
                                    run.font.name = style['font_name']
                                    run.font.color.rgb = style['color']
        print("OWQ part completed...")

        return doc

    # Usage:
    template_path = 'template.docx'
    mcq_file = 'chat_mcq.txt'
    owq_file = 'chat_owq.txt'

    filled_template = fill_template(template_path, mcq_file, owq_file)

    # Save the filled template to a new file
    output_path = 'output.docx'
    filled_template.save(output_path)

if __name__ == "__main__":
    generate_output()